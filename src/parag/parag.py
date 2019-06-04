
import os
import sys
import random

try:
    _path = os.path.split(os.path.split(os.path.abspath(__file__))[0])[0]
    sys.path.append(_path)
except:
    pass

_path = os.path.split(os.path.split(os.path.abspath("__file__"))[0])[0]
sys.path.append(_path)

from docx                         import Document
from PyQt5.QtCore                 import *
from PyQt5.QtGui                  import *
from PyQt5.QtWidgets              import * 
from parag_widgets.parag_widgets  import *
from parag_widgets.parag_css      import *
from parag_icons.parag_icons      import Parag_Icon
from parag_icons.parag_icons      import Parag_Pixmap
from datetime                     import datetime
from parag_db.parag_db            import PARAG_DB
from docx                         import Document
from datetime                     import datetime

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""

PARAG_NUMBER_OF_TEST_QUESTIONS = 10
PARAG_MIN_CORECT_QUESTIONS     = 8
PARAG_ANSWER_OPTIONS           = ["a","b","c","d","e","f","g","h"]
PARAG_QUESTIONS_HEADING        = """
INTREBARI EXAMEN ACORDARE/PRELUNGIRE
LICENŢĂ PILOT AERONAVE ULTRAUŞOARE
CLASA PARAPANTĂ
""" 

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Parag_UI(QMainWindow):

    def __init__(self,db):

        QMainWindow.__init__(self)

        self.stack_index  = 0
        self.db           = db      

        self.draw_gui() 

    def draw_gui(self):

        self.setWindowTitle("Licenta Parapanta")
        self.setWindowIcon(Parag_Icon("parag"))
        self.setMinimumSize(1300, 800)       
        self.setMinimumHeight(500)

        self.setStyleSheet(PARAG_CSS)

        self.wdg_central = QWidget()
        
        #selection tree
        self.tree = Parag_WDG_Tree()
        self.tree.setFixedWidth(260)
        self.populate_tree()
        self.tree.currentItemChanged.connect(self.tree_select)

        self.toolbar_layout = QHBoxLayout()

        #left tree layout area
        self.tree_area = QVBoxLayout()   
        self.tree_area.addLayout(self.toolbar_layout) 
        self.tree_area.addWidget(self.tree)   

        #context stack
        self.context = Parag_WDG_Stack()
        self.populate_stack()

        #Layout       
        self.top_layout = QHBoxLayout()
        self.top_layout.addLayout(self.tree_area)
        self.top_layout.addWidget(self.context)

        self.main_layout = QVBoxLayout()
        self.main_layout.addLayout(self.top_layout)

        self.wdg_central.setLayout(self.main_layout)  

        self.setCentralWidget(self.wdg_central)
        self.activateWindow() 

    def showEvent(self, event):    

        pass
        
    def tree_select(self, current, previous):

        self.context.setCurrentIndex(0) 

        _item_cfg = str(current.data(1,Qt.UserRole))

        #set context widget based on coresponding state     
        self.context.setCurrentIndex(int(_item_cfg))  

    def populate_tree(self):

        self.tree.clear()

        self.stack_index = 1 

        _parent = self.tree.invisibleRootItem()

        for _category in self.db:

            _item = QTreeWidgetItem(_parent)
            _item.setData(0, Qt.EditRole, _category.name)             

            _serial_cfg = self.stack_index                    
            _item.setData(1, Qt.UserRole, _serial_cfg)

            _item.setIcon(0,Parag_Icon("test"))

            self.stack_index += 1

    def populate_stack(self):

        self.context.insertWidget(0, QWidget())

        for _idx in range(len(self.db)):

            self.context.insertWidget(_idx + 1, Parag_WDG_Desktop(self.db[_idx])) 

        self.context.setCurrentIndex(0)  

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Parag_Model_Test(object):

    def __init__(self):

        self.questions = []

    def clear(self):

        for _question in self.questions:

            for _answer in _question.answers:

                _answer.is_selected = False

    def get_result(self):

        _total    = len(self.questions)
        _corect   = 0
        _incorect = 0

        for _question in self.questions:

            _is_corect = True

            for _answer in _question.answers:

                if _answer.is_corect:
                    if _answer.is_selected:
                        _is_corect = True
                    else:
                        _is_corect = False
                        break
                else:
                    if _answer.is_selected:
                        _is_corect = False
                        break
                    else:
                        _is_corect = True

            if _is_corect:
                _corect += 1
            else:
                _incorect += 1

        return _total,_corect,_incorect

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Parag_WDG_Desktop(QWidget):

    def __init__(self,category):

        QWidget.__init__(self)

        self.category = category

        self.draw_gui()

    def draw_gui(self):

        self.bt_test  = Parag_WDG_Button("Test",              Parag_Icon("test_normal"),   Parag_Icon("test_hover"),  "#606060")
        self.bt_learn = Parag_WDG_Button("Invata",            Parag_Icon("learn_normal"),  Parag_Icon("learn_hover"), "#606060")
        self.bt_gen   = Parag_WDG_Button("Genereaza Test",    Parag_Icon("generate_test"), Parag_Icon("generate_test"), "#606060")
        self.bt_exp   = Parag_WDG_Button("Exporta Intrebari", Parag_Icon("generate_doc"), Parag_Icon("generate_doc"), "#606060")

        self.bt_test.setIconSize(QSize(100,100))
        self.bt_learn.setIconSize(QSize(100,100))
        self.bt_gen.setIconSize(QSize(100,100))
        self.bt_exp.setIconSize(QSize(100,100))

        self.bt_test.clicked.connect(self.clbk_bt_test)
        self.bt_learn.clicked.connect(self.clbk_bt_learn)
        self.bt_gen.clicked.connect(self.clbk_bt_gen)
        self.bt_exp.clicked.connect(self.clbk_bt_exp)

        self.wdg_test = Parag_WDG_Desktop_Test(self,self.category)

        self.bt_layout = QHBoxLayout()
        self.bt_layout.addWidget(self.bt_learn)
        self.bt_layout.addWidget(self.bt_test)
        self.bt_layout.addWidget(self.bt_gen)
        self.bt_layout.addWidget(self.bt_exp)

        self.main_layout = QVBoxLayout()
        self.main_layout.addLayout(self.bt_layout)
        self.main_layout.addWidget(self.wdg_test)

        self.wdg_test.hide()

        self.setLayout(self.main_layout)  

    def get_docs_path(self):

        _path = os.path.abspath(__file__)
        _path = os.path.split(_path)[0]
        _path = os.path.join(_path,"docs")

        if not os.path.exists(_path):

            os.mkdir(_path)

        return _path

    def get_generated_doc_path(self,name):

        _path = self.get_docs_path()

        _timestamp = str(datetime.now())

        _file_name = "intrebari_%s_%s.docx" % (name,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

        _path = os.path.join(_path,_file_name)

        return _path,_timestamp

    def clbk_bt_test(self,state):

        self.wdg_test.show()

        self.bt_test.hide()

        self.bt_learn.hide()

        self.bt_gen.hide()

        self.bt_exp.hide()

        self.wdg_test.start("test")

    def clbk_bt_learn(self,state):

        self.wdg_test.show()

        self.bt_test.hide()

        self.bt_learn.hide()

        self.bt_gen.hide()

        self.bt_exp.hide()

        self.wdg_test.start("learn")

    def clbk_bt_gen(self,state):

        pass

        #TODO

    def clbk_bt_exp(self,state):

        _path, _timestamp = self.get_generated_doc_path(self.category.name)

        _document = Document()

        _paragraph = _document.add_paragraph()
        _paragraph.add_run(PARAG_QUESTIONS_HEADING).bold = True
        _document.add_paragraph("")
        _document.add_paragraph("")

        _question_count = 1

        for _question in self.category.questions:

            _question_text = "%s. %s" % (_question_count,_question.text)

            _question_count += 1

            _paragraph = _document.add_paragraph(_question_text)

            #TODO add images

            _answer_count = 1

            for _answer in _question.answers:

                _answer_text = "    %s) %s" % (PARAG_ANSWER_OPTIONS[_answer_count],_answer.text)

                _answer_count += 1

                _paragraph = _document.add_paragraph()
                _run       = _paragraph.add_run(_answer_text)
                _run.bold  = _answer.corect

            _document.add_paragraph("")

        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("Generat la data: %s" % (_timestamp,))


        _document.save(_path)

        _msg = QMessageBox()
        _msg.setWindowIcon(Parag_Icon("parag"))
        _msg.setIcon(QMessageBox.Information)
        _msg.setText("Generat intrebari in fisierul\n%s" % (os.path.split(_path)[1],))
        _msg.setWindowTitle("Generat Intrebari %s" % (self.category.name.upper(),))
        _msg.exec_()

        os.startfile(_path)

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Parag_WDG_Desktop_Test(QWidget):

    def __init__(self,parent,category):

        QWidget.__init__(self)

        self.category        = category
        self.question_number = 0
        self.parent          = parent
        self.test_type       = ""
        self.test_learn      = Parag_Model_Test()
        self.test_exam       = Parag_Model_Test()

        self.test_learn.questions = self.category.questions

        self.draw_gui()

    def draw_gui(self):

        self.bt_next   = Parag_WDG_Small_Button( Parag_Icon("next_normal"),       Parag_Icon("next_hover"),     self.clbk_next , "Intrebare urmatoare")
        self.bt_prev   = Parag_WDG_Small_Button( Parag_Icon("previous_normal"),   Parag_Icon("previous_hover"), self.clbk_prev, "Intrebarea precedenta")
        self.bt_result = Parag_WDG_Small_Button( Parag_Icon("results_normal"),    Parag_Icon("result_hover"),   self.clbk_result, "Rezultat")
        self.bt_close  = Parag_WDG_Small_Button( Parag_Icon("close_normal"),      Parag_Icon("close_hover"),    self.clbk_close, "Inchide")

        self.bt_next.setIconSize(QSize(50,50))
        self.bt_prev.setIconSize(QSize(50,50))
        self.bt_result.setIconSize(QSize(50,50))
        self.bt_close.setIconSize(QSize(50,50))

        self.lbl_status = Parag_WDG_Label()
        self.lbl_result = Parag_WDG_Label()

        self.bt_layout = QHBoxLayout()
        self.bt_layout.addWidget(self.bt_prev)
        self.bt_layout.addWidget(self.bt_close)
        self.bt_layout.addWidget(self.bt_result)
        self.bt_layout.addWidget(self.bt_next)

        self.wdg_question = Parag_WDG_Question()

        self.main_layout = QVBoxLayout()
        
        self.main_layout.addLayout(self.bt_layout)
        self.main_layout.addWidget(self.wdg_question)
        self.main_layout.addWidget(self.lbl_status)
        self.main_layout.addWidget(self.lbl_result)

        self.setLayout(self.main_layout) 

        self.lbl_result.hide()

        self.bt_next.hide()
        self.bt_prev.hide()
        self.bt_close.hide()
        self.bt_result.hide()

    def start(self,test_type):

        self.test_type = test_type

        self.question_number = 0

        self.wdg_question.show()

        if self.test_type == "learn":            
            self.test_learn.clear()
            self.wdg_question.populate(self.test_learn.questions[self.question_number])
            self.bt_next.show()
            self.bt_close.show()
            self.bt_result.show()
            self.bt_prev.show()
        else:            
            self.test_exam.clear()
            self.bt_next.show()
            self.bt_prev.show()
            self.bt_result.show()
            self.get_test_questions()
            self.wdg_question.populate(self.test_exam.questions[self.question_number])

        self.lbl_result.hide()
        self.set_status()

    def get_test_questions(self):

        global PARAG_NUMBER_OF_TEST_QUESTIONS

        if PARAG_NUMBER_OF_TEST_QUESTIONS < len(self.test_learn.questions):

            _questions_indexes = random.sample(range(len(self.test_learn.questions)), PARAG_NUMBER_OF_TEST_QUESTIONS)

            self.test_exam.questions = [self.test_learn.questions[_index] for _index in _questions_indexes]
        else:
            self.test_exam.questions = self.test_learn.questions

    def clbk_next(self,state):

        if not self.is_end(self.question_number):

            self.question_number += 1

            if self.test_type == "learn":
                self.wdg_question.populate(self.test_learn.questions[self.question_number])
            else:
                self.wdg_question.populate(self.test_exam.questions[self.question_number])

            self.set_status()
        else:
            pass
            #TODO complete test

    def clbk_prev(self,state):

        if not self.is_begining(self.question_number):

            self.question_number -= 1

            if self.test_type == "learn":
                self.wdg_question.populate(self.test_learn.questions[self.question_number])
            else:
                self.wdg_question.populate(self.test_exam.questions[self.question_number])

            self.set_status()

        else:

            pass

    def clbk_close(self,state):

        self.parent.bt_test.show()
        self.parent.bt_learn.show()
        self.parent.bt_gen.show()
        self.parent.bt_exp.show()
        self.parent.wdg_test.hide()

    def clbk_result(self,state):

        _corect   = 0
        _incorect = 0
        _total    = 0

        if self.test_type == "learn":

            if self.test_learn.questions[self.question_number].answers[0].corect:
                self.wdg_question.rd_answer_a.setStyleSheet("QCheckBox { color: green }")
            else:
                self.wdg_question.rd_answer_a.setStyleSheet("QCheckBox { color: red }")

            if self.test_learn.questions[self.question_number].answers[1].corect:
                self.wdg_question.rd_answer_b.setStyleSheet("QCheckBox { color: green }")
            else:
                self.wdg_question.rd_answer_b.setStyleSheet("QCheckBox { color: red }")

            if self.test_learn.questions[self.question_number].answers[2].corect:
                self.wdg_question.rd_answer_c.setStyleSheet("QCheckBox { color: green }")
            else:
                self.wdg_question.rd_answer_c.setStyleSheet("QCheckBox { color: red }")

        else:
            _total,_corect,_incorect = self.test_exam.get_result()

            self.lbl_status.setText("Intrebari[%s] Corecte[%s] Incorecte[%s]" % (_total,_corect,_incorect))

            self.bt_next.hide()
            self.bt_prev.hide()
            self.bt_result.hide()
            self.bt_close.show()
            self.wdg_question.hide()

            self.lbl_result.show()

            if _corect > PARAG_MIN_CORECT_QUESTIONS:
                self.lbl_result.setText("ADMIS")
                self.lbl_result.setStyleSheet("QLabel { background-color : #14c941; font: 18pt; color: #ffffff}")
                self.lbl_result.setAlignment(Qt.AlignCenter)
            else:
                self.lbl_result.setText("PICAT")
                self.lbl_result.setStyleSheet("QLabel { background-color : #ba2012; font: 18pt;  color: #ffffff}")
                self.lbl_result.setAlignment(Qt.AlignCenter)

    def set_status(self):

        if self.test_type == "learn":
            self.lbl_status.setText("Intrebarea[%s/%s] " % (self.question_number + 1,len(self.test_learn.questions)))
        else:
            self.lbl_status.setText("Intrebarea[%s/%s]" % (self.question_number + 1,len(self.test_exam.questions)))

    def is_end(self,question_number):

        _state = False

        if self.test_type == "learn":
            _state = (question_number + 1) >= len(self.test_learn.questions)
        else:
            _state = (question_number + 1) >= len(self.test_exam.questions)

        return _state

    def is_begining(self,question_number):

        return question_number <= 0

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Parag_WDG_Question(QWidget):

    def __init__(self):

        QWidget.__init__(self)

        self.draw_gui()

        self.question = None

    def draw_gui(self):

        self.lbl_question = Parag_WDG_Label()
        self.rd_answer_a  = Parag_WDG_CheckBox("")
        self.rd_answer_b  = Parag_WDG_CheckBox("")
        self.rd_answer_c  = Parag_WDG_CheckBox("")


        self.rd_answer_a.stateChanged.connect(self.clbk_answer_a)
        self.rd_answer_b.stateChanged.connect(self.clbk_answer_b)
        self.rd_answer_c.stateChanged.connect(self.clbk_answer_c)

        self.main_layout = QVBoxLayout()
        self.main_layout.addWidget(self.lbl_question)
        self.main_layout.addWidget(self.rd_answer_a)
        self.main_layout.addWidget(self.rd_answer_b)
        self.main_layout.addWidget(self.rd_answer_c)

        self.setLayout(self.main_layout)

        self.lbl_question.hide()
        self.rd_answer_a.hide()
        self.rd_answer_b.hide()
        self.rd_answer_c.hide()

    def populate(self,question):

        self.rd_answer_a.setStyleSheet("QCheckBox { color: #b1b1b1 }")
        self.rd_answer_b.setStyleSheet("QCheckBox { color: #b1b1b1 }")
        self.rd_answer_c.setStyleSheet("QCheckBox { color: #b1b1b1 }")


        self.question = question

        self.lbl_question.show()
        self.rd_answer_a.show()
        self.rd_answer_b.show()
        self.rd_answer_c.show()

        self.lbl_question.setText(self.question.text)
        self.rd_answer_a.setText(self.question.answers[0].text)
        self.rd_answer_b.setText(self.question.answers[1].text)
        self.rd_answer_c.setText(self.question.answers[2].text)

        if self.question.answers[0].selected:
            self.rd_answer_a.setCheckState(Qt.Checked)
        else:
            self.rd_answer_a.setCheckState(Qt.Unchecked)

        if self.question.answers[1].selected:
            self.rd_answer_b.setCheckState(Qt.Checked)
        else:
            self.rd_answer_b.setCheckState(Qt.Unchecked)

        if self.question.answers[2].selected:
            self.rd_answer_c.setCheckState(Qt.Checked)
        else:
            self.rd_answer_c.setCheckState(Qt.Unchecked)

    def clbk_answer_a(self,state):
        
        self.question.answers[0].selected = state == Qt.Checked

    def clbk_answer_b(self,state):
        
        self.question.answers[1].selected = state == Qt.Checked

    def clbk_answer_c(self,state):
        
        self.question.answers[2].selected = state == Qt.Checked

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Parag(object):

    def __init__(self):

        global PARAG_DB

        self.app = QApplication(sys.argv)  

        _ui  = Parag_UI(PARAG_DB)   

        _ui.show()

        sys.exit(self.app.exec_())

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
if __name__ == "__main__":

    Parag()

